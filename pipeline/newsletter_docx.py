"""Word document renderer for the FMCG Deal newsletter."""

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from config import OUTPUT_DIR

_DEAL_EMOJI = {
    "acquisition": "Acquisition",
    "merger": "Merger",
    "jv": "Joint Venture",
    "investment": "Investment",
    "divestiture": "Divestiture",
    "ipo": "IPO",
    "partnership": "Partnership",
    "other": "Other",
}

_STATUS_LABEL = {
    "completed": "Completed",
    "announced": "Announced",
    "rumored": "Rumored",
    "in-progress": "In Progress",
}


def _set_cell_shading(cell, color_hex: str):
    """Apply background shading to a table cell."""
    from docx.oxml.ns import qn
    from lxml import etree
    shading = etree.SubElement(cell._element.get_or_add_tcPr(), qn("w:shd"))
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")


def _add_styled_heading(doc: Document, text: str, level: int = 1):
    """Add a heading with custom styling."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return heading


def _add_bullet_point(doc: Document, text: str):
    """Add a bullet point paragraph."""
    p = doc.add_paragraph(text, style="List Bullet")
    return p


def _has_structured_fields(deal: dict) -> bool:
    return deal.get("deal_type") is not None


def render_docx(sections: dict, deals: list, run_date: str) -> str:
    """Render newsletter as a Word document.

    Returns the path to the generated .docx file.
    """
    doc = Document()

    # -- Style tweaks --
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Title
    title = doc.add_heading(f"FMCG Deal Pulse — Week of {run_date}", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    # ---------- TL;DR ----------
    _add_styled_heading(doc, "TL;DR", level=1)
    summary_text = sections.get("executive_summary", "")
    for line in summary_text.strip().splitlines():
        line = line.strip().lstrip("- ").strip()
        if line:
            _add_bullet_point(doc, line)

    # ---------- Deal of the Week ----------
    _add_styled_heading(doc, "Deal of the Week", level=1)

    if deals:
        top = deals[0]
        if _has_structured_fields(top):
            # Metadata table
            table = doc.add_table(rows=1, cols=4)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            hdr[0].text = "Type"
            hdr[1].text = "Value"
            hdr[2].text = "Sector"
            hdr[3].text = "Status"

            row = table.add_row().cells
            row[0].text = _DEAL_EMOJI.get(top.get("deal_type", ""), "Other")
            row[1].text = top.get("deal_value_structured") or "Undisclosed"
            row[2].text = top.get("sector") or "FMCG"
            row[3].text = _STATUS_LABEL.get(top.get("deal_status", ""), top.get("deal_status", ""))

            doc.add_paragraph()  # spacer

    # Narrative
    headline_text = sections.get("headline_deal", "")
    for para in headline_text.split("\n\n"):
        para = para.strip()
        if para:
            # Strip markdown bold markers for Word
            para = para.replace("**", "")
            doc.add_paragraph(para)

    # ---------- Deal Briefs ----------
    _add_styled_heading(doc, "Deal Briefs", level=1)

    briefs_pool = deals[1:6]
    if briefs_pool and any(_has_structured_fields(d) for d in briefs_pool):
        table = doc.add_table(rows=1, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Acquirer"
        hdr[1].text = "Target"
        hdr[2].text = "Type"
        hdr[3].text = "Value"
        hdr[4].text = "Sector"

        for d in briefs_pool:
            if _has_structured_fields(d):
                row = table.add_row().cells
                row[0].text = d.get("acquirer") or "Undisclosed"
                row[1].text = d.get("target") or "Undisclosed"
                row[2].text = _DEAL_EMOJI.get(d.get("deal_type", ""), "Other")
                row[3].text = d.get("deal_value_structured") or "Undisclosed"
                row[4].text = d.get("sector") or "FMCG"
    else:
        doc.add_paragraph(sections.get("deal_briefs", "No additional deals.").replace("**", ""))

    # ---------- Sector Pulse ----------
    _add_styled_heading(doc, "Sector Pulse", level=1)
    pulse_text = sections.get("sector_pulse", "")
    for para in pulse_text.split("\n\n"):
        para = para.strip().replace("**", "")
        if para:
            doc.add_paragraph(para)

    # ---------- Watchlist ----------
    _add_styled_heading(doc, "Watchlist", level=1)
    watch_deals = [d for d in deals if d.get("deal_status") in ("rumored", "in-progress")]
    if watch_deals:
        for d in watch_deals:
            status = _STATUS_LABEL.get(d.get("deal_status", ""), d.get("deal_status", ""))
            if _has_structured_fields(d):
                text = (
                    f"{d.get('acquirer', '?')} + {d.get('target', '?')} — "
                    f"{d.get('deal_value_structured', 'Undisclosed')} "
                    f"{_DEAL_EMOJI.get(d.get('deal_type', ''), 'deal')} ({status})"
                )
            else:
                text = f"{d.get('title', 'Untitled')} — {status}"
            _add_bullet_point(doc, text)
    else:
        doc.add_paragraph("No early-stage deals or rumors on the radar this week.")

    # ---------- Footer ----------
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(
        f"Generated by FMCG Deal Intelligence Pipeline — {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # Save
    OUTPUT_DIR.mkdir(exist_ok=True)
    docx_path = OUTPUT_DIR / "newsletter.docx"
    doc.save(str(docx_path))
    print(f"  [newsletter] Saved → {docx_path}")
    return str(docx_path)
